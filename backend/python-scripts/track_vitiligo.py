import sys
import json
import cv2
import numpy as np
import base64
from docx import Document
from docx.shared import Inches
import os
import traceback
import pickle
import warnings
from PIL import Image

DEBUG = True

def log_debug(message):
    if DEBUG:
        print(f"[DEBUG] {message}", file=sys.stderr)
        sys.stderr.flush()

def base64_to_image(base64_str):
    log_debug("Starting base64_to_image conversion")
    try:
        img_data = base64.b64decode(base64_str)
        np_arr = np.frombuffer(img_data, np.uint8)
        img = cv2.imdecode(np_arr, cv2.IMREAD_COLOR)
        log_debug(f"Successfully converted base64 to image. Shape: {img.shape}")
        return img
    except Exception as e:
        log_debug(f"Error in base64_to_image: {str(e)}")
        raise

def numpy_to_pil(image_np):
    log_debug("Converting numpy array to PIL image")
    try:
        pil_img = Image.fromarray(cv2.cvtColor(image_np, cv2.COLOR_BGR2RGB))
        log_debug("Successfully converted numpy to PIL")
        return pil_img
    except Exception as e:
        log_debug(f"Error in numpy_to_pil: {str(e)}")
        raise

def to_grayscale(image):
    return cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)

def segment_and_draw_contours(gray_img):
    _, thresh = cv2.threshold(gray_img, 180, 255, cv2.THRESH_BINARY)
    contours, _ = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    mask = cv2.cvtColor(gray_img, cv2.COLOR_GRAY2BGR)
    cv2.drawContours(mask, contours, -1, (0, 0, 255), 2)
    area = sum(cv2.contourArea(c) for c in contours)
    return mask, float(area)

def get_non_black_area(gray_img):
    _, black_mask = cv2.threshold(gray_img, 50, 255, cv2.THRESH_BINARY_INV)
    contours, _ = cv2.findContours(black_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    black_area = sum(cv2.contourArea(c) for c in contours)
    total_area = float(gray_img.shape[0] * gray_img.shape[1])
    return total_area - float(black_area)

def calculate_speed_rate(before_area, after_area, weeks):
    try:
        weeks = float(weeks)
        before_area = float(before_area)
        after_area = float(after_area)
        
        if before_area > 0 and weeks > 0:
            return ((after_area - before_area) / before_area) * 100 / weeks
        return 0.0
    except (ValueError, TypeError) as e:
        log_debug(f"Error in calculate_speed_rate: {str(e)}")
        return 0.0

def load_model(model_path=None):
    try:
        if model_path is None:
            script_dir = os.path.dirname(os.path.abspath(__file__))
            model_path = os.path.join(script_dir, 'viti.pkl')
        
        log_debug(f"Attempting to load model from: {model_path}")
        
        if not os.path.exists(model_path):
            raise FileNotFoundError(f"Model file not found at: {model_path}")
            
        with open(model_path, 'rb') as file:
            model_bundle = pickle.load(file)
            
        # Verify model bundle structure
        if 'model' not in model_bundle or 'label_encoder_treatment' not in model_bundle:
            raise ValueError("Model bundle is missing required components")
            
        log_debug("Model loaded successfully")
        return model_bundle
        
    except Exception as e:
        log_debug(f"Model loading failed: {str(e)}")
        raise

def generate_report(name, age, gender, weeks, before_image, after_image, 
                   before_segmented, after_segmented, change_percentage, treatment_recommendation):
    log_debug("Starting report generation")
    try:
        doc = Document()
        doc.add_heading('Vitiligo Progress Report', 0)

        # Patient Information
        doc.add_heading('Patient Information', level=1)
        doc.add_paragraph(f"Name: {name}")
        doc.add_paragraph(f"Age: {int(age)}")
        doc.add_paragraph(f"Gender: {gender}")
        doc.add_paragraph(f"Duration Between Photos: {float(weeks):.1f} weeks")

        # Treatment Recommendation
        doc.add_heading('Treatment Recommendation', level=1)
        doc.add_paragraph(treatment_recommendation)

        # Progress Summary
        doc.add_heading('Progress Summary', level=1)
        doc.add_paragraph(f"Change in Affected Area: {float(change_percentage):.2f}%")

        # Create reports directory if it doesn't exist
        script_dir = os.path.dirname(os.path.abspath(__file__))
        reports_dir = os.path.join(script_dir, 'generated_reports')
        os.makedirs(reports_dir, exist_ok=True)

        # Before Treatment Images
        doc.add_heading('Before Treatment', level=2)
        
        before_temp = os.path.join(reports_dir, 'before_temp.png')
        numpy_to_pil(before_image).save(before_temp)
        doc.add_paragraph("Original Image:")
        doc.add_picture(before_temp, width=Inches(3))
        
        before_seg_temp = os.path.join(reports_dir, 'before_seg_temp.png')
        numpy_to_pil(before_segmented).save(before_seg_temp)
        doc.add_paragraph("Segmented Image (affected areas in red):")
        doc.add_picture(before_seg_temp, width=Inches(3))

        # After Treatment Images
        doc.add_heading('After Treatment', level=2)
        after_temp = os.path.join(reports_dir, 'after_temp.png')
        numpy_to_pil(after_image).save(after_temp)
        doc.add_paragraph("Original Image:")
        doc.add_picture(after_temp, width=Inches(3))
        
        after_seg_temp = os.path.join(reports_dir, 'after_seg_temp.png')
        numpy_to_pil(after_segmented).save(after_seg_temp)
        doc.add_paragraph("Segmented Image (affected areas in red):")
        doc.add_picture(after_seg_temp, width=Inches(3))

        # Clean up temporary files
        for temp_file in [before_temp, before_seg_temp, after_temp, after_seg_temp]:
            if os.path.exists(temp_file):
                os.remove(temp_file)

        # Save the final report
        report_path = os.path.join(reports_dir, f"{name.replace(' ', '_')}_vitiligo_report.docx")
        doc.save(report_path)
        
        log_debug(f"Report successfully generated at: {report_path}")
        return report_path

    except Exception as e:
        log_debug(f"Error in generate_report: {str(e)}")
        raise

def main():
    try:
        warnings.filterwarnings("ignore", category=UserWarning)
        
        log_debug("\n==== Starting Vitiligo Analysis ====")
        input_data = sys.stdin.read()
        data = json.loads(input_data)
        
        # Convert all numeric inputs to proper types
        data['age'] = int(data.get('age', 0))
        data['weeks'] = float(data.get('weeks', 0))
        
        log_debug(f"Processing data for patient: {data['name']}")

        # Process images
        before_image = base64_to_image(data['before_image'])
        after_image = base64_to_image(data['after_image'])
        
        # Convert to grayscale
        before_gray = to_grayscale(before_image)
        after_gray = to_grayscale(after_image)

        # Segment images and calculate areas
        before_mask, before_area = segment_and_draw_contours(before_gray)
        after_mask, after_area = segment_and_draw_contours(after_gray)
        
        # Calculate valid skin area
        valid_area = get_non_black_area(after_gray)
        
        # Calculate percentages and speed rate
        percent_before = (float(before_area) / float(valid_area)) * 100 if valid_area > 0 else 0.0
        percent_after = (float(after_area) / float(valid_area)) * 100 if valid_area > 0 else 0.0
        change_percentage = ((float(after_area) - float(before_area)) / float(before_area)) * 100 if before_area > 0 else 0.0
        speed_rate = calculate_speed_rate(before_area, after_area, data['weeks'])

        # Load model and make prediction
        model_bundle = load_model()
        features = np.array([[
            float(data['age']), 
            float(data['weeks']), 
            float(speed_rate), 
            float(percent_before), 
            float(percent_after)
        ]])
        
        predicted_label = model_bundle['model'].predict(features)[0]
        treatment_recommendation = model_bundle['label_encoder_treatment'].inverse_transform([predicted_label])[0]

        # Generate report
        report_path = generate_report(
            data['name'], 
            data['age'], 
            data.get('gender', ''), 
            data['weeks'],
            before_image, 
            after_image, 
            before_mask, 
            after_mask, 
            change_percentage, 
            treatment_recommendation
        )

        # Return results
        print(json.dumps({
            "status": "success",
            "report_path": report_path,
            "before_area": float(before_area),
            "after_area": float(after_area),
            "change_percentage": float(change_percentage),
            "treatment_recommendation": treatment_recommendation,
            "speed_rate": float(speed_rate)
        }))

    except Exception as e:
        log_debug(f"\n==== ERROR OCCURRED ====\n{traceback.format_exc()}")
        print(json.dumps({
            "status": "error",
            "message": str(e),
            "traceback": traceback.format_exc()
        }))

if __name__ == "__main__":
    main()